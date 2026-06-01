import sqlite3
import os
import re
import docx
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Paths
DB_PATH = r"C:\Users\wagahsan\OneDrive - Minnesota State\Desktop\Shadman Ahsan\TRC-AI-Assistant\trc_ai.db"
OUTPUT_DIR = r"C:\Users\wagahsan\OneDrive - Minnesota State\Desktop\Shadman Ahsan\TRC-AI-Assistant\Collected_KB_Documents"

def clean_filename(filename):
    # Remove invalid filename characters on Windows
    return re.sub(r'[\\/*?:"<>|]', "", filename).strip()

def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    """
    Adds a hyperlink to a paragraph.
    """
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)

    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    new_run.append(rPr)
    text_node = OxmlElement('w:t')
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def convert_articles():
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)
        print(f"Created output directory: {OUTPUT_DIR}")

    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    
    cur.execute("SELECT id, content FROM kb WHERE source = 'SMSU KB Scraper'")
    rows = cur.fetchall()
    
    print(f"Found {len(rows)} articles to convert...")
    
    import docx # import here to use REL relationships in add_hyperlink
    
    success_count = 0
    for id_, raw_content in rows:
        try:
            # Parse article parts using regex/string processing
            title_match = re.match(r"\*\*(.*?)\*\*", raw_content)
            title = title_match.group(1).strip() if title_match else f"Article_{id_}"
            
            url_match = re.search(r"🔗 \*\*Official KB Article\*\*: (https://\S+)", raw_content)
            url = url_match.group(1).strip() if url_match else ""
            
            cat_match = re.search(r"Category: (.*?)<br>", raw_content)
            category = cat_match.group(1).strip() if cat_match else "General"
            
            tags_match = re.search(r"Tags: (.*?)<br>", raw_content)
            tags = tags_match.group(1).strip() if tags_match else "None"
            
            # Extract actual body text
            # Find the index of the end of the metadata block (e.g. after <br><br> or tags)
            body_start = raw_content.find("<br><br>")
            if body_start != -1:
                body_content = raw_content[body_start + 8:].strip()
            else:
                body_content = raw_content
                
            # Create a professional Word document
            doc = Document()
            
            # Document Title
            h1 = doc.add_heading(title, level=1)
            h1.runs[0].font.name = 'Inter'
            h1.runs[0].font.size = Pt(20)
            h1.runs[0].font.bold = True
            
            # Metadata Box
            p_meta = doc.add_paragraph()
            p_meta.paragraph_format.space_before = Pt(12)
            p_meta.paragraph_format.space_after = Pt(18)
            
            run_url_lbl = p_meta.add_run("Official URL: ")
            run_url_lbl.bold = True
            run_url_lbl.font.size = Pt(10)
            if url:
                add_hyperlink(p_meta, url, url)
            else:
                p_meta.add_run("Not Available")
                
            p_meta.add_run("\n")
            run_cat = p_meta.add_run("Category: ")
            run_cat.bold = True
            run_cat.font.size = Pt(10)
            p_meta.add_run(category)
            
            p_meta.add_run("\n")
            run_tags = p_meta.add_run("Tags: ")
            run_tags.bold = True
            run_tags.font.size = Pt(10)
            p_meta.add_run(tags)
            
            # Horizontal line separator
            doc.add_paragraph("_____________________________________________________\n")
            
            # Body content formatting
            # Process body text paragraph by paragraph
            paragraphs = body_content.split('\n')
            for p_text in paragraphs:
                p_text = p_text.strip()
                if not p_text:
                    continue
                
                # Check for headings in raw content
                if p_text.startswith('###'):
                    h = doc.add_heading(p_text.replace('###', '').strip(), level=3)
                    h.runs[0].font.name = 'Inter'
                    h.runs[0].font.size = Pt(12)
                    h.runs[0].font.bold = True
                elif p_text.startswith('##'):
                    h = doc.add_heading(p_text.replace('##', '').strip(), level=2)
                    h.runs[0].font.name = 'Inter'
                    h.runs[0].font.size = Pt(14)
                    h.runs[0].font.bold = True
                elif p_text.startswith('#'):
                    h = doc.add_heading(p_text.replace('#', '').strip(), level=1)
                    h.runs[0].font.name = 'Inter'
                    h.runs[0].font.size = Pt(16)
                    h.runs[0].font.bold = True
                else:
                    p = doc.add_paragraph(p_text)
                    p.paragraph_format.line_spacing = 1.15
                    p.paragraph_format.space_after = Pt(6)
            
            # Add accessibility notes at bottom if needed
            doc.add_paragraph("\n\nAccessibility Review & Update Log:")
            doc.add_paragraph("- [ ] Content accuracy verified\n- [ ] Color contrast checked\n- [ ] Image alt-text verified\n- [ ] Document structure and headings structured")
            
            # Save document
            safe_title = clean_filename(title)
            # Limit length of title to avoid path-too-long issues on Windows
            if len(safe_title) > 80:
                safe_title = safe_title[:80]
                
            file_name = f"{safe_title}.docx"
            file_path = os.path.join(OUTPUT_DIR, file_name)
            
            # Handle duplicate file names safely
            counter = 1
            while os.path.exists(file_path):
                file_name = f"{safe_title}_{counter}.docx"
                file_path = os.path.join(OUTPUT_DIR, file_name)
                counter += 1
                
            doc.save(file_path)
            success_count += 1
            
        except Exception as e:
            print(f"Error converting article {id_}: {e}")
            
    print(f"\nSuccessfully converted {success_count} / {len(rows)} articles to Word documents!")
    conn.close()

if __name__ == "__main__":
    convert_articles()
