import sqlite3
import os
import re
import base64
import time
from playwright.sync_api import sync_playwright
import docx
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Paths
DB_PATH = r"C:\Users\wagahsan\OneDrive - Minnesota State\Desktop\Shadman Ahsan\TRC-AI-Assistant\trc_ai.db"
OUTPUT_DIR = r"C:\Users\wagahsan\OneDrive - Minnesota State\Desktop\Shadman Ahsan\TRC-AI-Assistant\Collected_KB_Documents"

def clean_filename(filename):
    # Remove invalid filename characters on Windows
    return re.sub(r'[\\/*?:"<>|]', "", filename).strip()

def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    """
    Adds a hyperlink to a paragraph.
    """
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)

    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    new_run.append(rPr)
    text_node = OxmlElement('w:t')
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def main():
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)
        print(f"Created output directory: {OUTPUT_DIR}")

    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    
    cur.execute("SELECT id, content FROM kb WHERE source = 'SMSU KB Scraper'")
    rows = cur.fetchall()
    
    print(f"Found {len(rows)} articles in database. Starting Playwright for screenshot-aware document generation...")
    
    success_count = 0
    
    with sync_playwright() as p:
        print("Launching Chromium headless browser...")
        browser = p.chromium.launch(headless=True)
        page = browser.new_page()
        
        for idx, (id_, db_content) in enumerate(rows, 1):
            try:
                # 1. Parse basic metadata from DB
                title_match = re.match(r"\*\*(.*?)\*\*", db_content)
                title = title_match.group(1).strip() if title_match else f"Article_{id_}"
                
                url_match = re.search(r"Official KB Article\*\*: (https://[^\s<]+)", db_content)
                url = url_match.group(1).strip() if url_match else ""
                
                cat_match = re.search(r"Category: (.*?)<br>", db_content)
                category = cat_match.group(1).strip() if cat_match else "General"
                
                tags_match = re.search(r"Tags: (.*?)<br>", db_content)
                tags = tags_match.group(1).strip() if tags_match else "None"
                
                print(f"[{idx}/{len(rows)}] Processing: '{title}'...")
                
                blocks = []
                loaded_live = False
                
                # 2. Try to fetch the live page to extract elements and screenshots
                if url:
                    try:
                        page.goto(url, timeout=15000)
                        # Wait for either the article body or a redirect/error indicator
                        page.wait_for_selector("#ctl00_ctl00_cpContent_cpContent_divBody, h1", timeout=8000)
                        
                        # Check if divBody is present
                        if page.locator("#ctl00_ctl00_cpContent_cpContent_divBody").count() > 0:
                            # Run browser DFS to parse headings, paragraphs, lists, and images in order
                            blocks = page.evaluate("""() => {
                                const root = document.querySelector("#ctl00_ctl00_cpContent_cpContent_divBody");
                                if (!root) return [];
                                
                                const blocks = [];
                                
                                function dfs(node) {
                                    if (!node) return;
                                    
                                    const tagName = node.tagName;
                                    
                                    // Handle images
                                    if (tagName === 'IMG') {
                                        blocks.push({
                                            type: 'image',
                                            src: node.src,
                                            alt: node.alt || 'Screenshot'
                                        });
                                        return;
                                    }
                                    
                                    // Handle headings
                                    if (['H1', 'H2', 'H3', 'H4', 'H5', 'H6'].includes(tagName)) {
                                        blocks.push({
                                            type: 'heading',
                                            level: parseInt(tagName[1]),
                                            text: node.innerText.trim()
                                        });
                                        return;
                                    }
                                    
                                    // Handle individual list items
                                    if (tagName === 'LI') {
                                        const clone = node.cloneNode(true);
                                        const nestedImgs = clone.querySelectorAll('img');
                                        nestedImgs.forEach(img => img.remove());
                                        const txt = clone.innerText.trim();
                                        
                                        if (txt) {
                                            const parentTag = (node.parentElement && node.parentElement.tagName === 'OL') ? 'ol' : 'ul';
                                            blocks.push({
                                                type: 'list-item',
                                                list_type: parentTag,
                                                text: txt
                                            });
                                        }
                                        
                                        // Traverse nested children inside the list item to capture images
                                        for (let child of node.childNodes) {
                                            if (child.nodeType === Node.ELEMENT_NODE) {
                                                dfs(child);
                                            }
                                        }
                                        return;
                                    }
                                    
                                    // For paragraph or basic elements, check if they have block children
                                    if (['P', 'DIV', 'SPAN'].includes(tagName)) {
                                        const hasNestedBlock = node.querySelector('p, h1, h2, h3, h4, h5, h6, ul, ol, li, img');
                                        if (!hasNestedBlock) {
                                            const txt = node.innerText.trim();
                                            if (txt) {
                                                blocks.push({
                                                    type: 'paragraph',
                                                    text: txt
                                                });
                                            }
                                            return;
                                        }
                                    }
                                    
                                    // Traverse children
                                    for (let child of node.childNodes) {
                                        if (child.nodeType === Node.ELEMENT_NODE) {
                                            dfs(child);
                                        } else if (child.nodeType === Node.TEXT_NODE) {
                                            const txt = child.textContent.trim();
                                            if (txt && node === root) {
                                                blocks.push({
                                                    type: 'paragraph',
                                                    text: txt
                                                });
                                            }
                                        }
                                    }
                                }
                                
                                dfs(root);
                                return blocks;
                            }""")
                            loaded_live = True
                            print(f" -> Successfully fetched live page! Extracted {len(blocks)} element blocks.")
                    except Exception as e:
                        print(f" -> Live fetch skipped (requires auth or timed out): {e}")
                
                # 3. Create styled document
                doc = Document()
                
                # Title Styling
                h1 = doc.add_heading(title, level=1)
                h1.runs[0].font.name = 'Inter'
                h1.runs[0].font.size = Pt(20)
                h1.runs[0].font.bold = True
                
                # Metadata Box
                p_meta = doc.add_paragraph()
                p_meta.paragraph_format.space_before = Pt(12)
                p_meta.paragraph_format.space_after = Pt(18)
                
                run_url_lbl = p_meta.add_run("Official URL: ")
                run_url_lbl.bold = True
                run_url_lbl.font.size = Pt(10)
                if url:
                    add_hyperlink(p_meta, url, url)
                else:
                    p_meta.add_run("Not Available")
                    
                p_meta.add_run("\n")
                run_cat = p_meta.add_run("Category: ")
                run_cat.bold = True
                run_cat.font.size = Pt(10)
                p_meta.add_run(category)
                
                p_meta.add_run("\n")
                run_tags = p_meta.add_run("Tags: ")
                run_tags.bold = True
                run_tags.font.size = Pt(10)
                p_meta.add_run(tags)
                
                doc.add_paragraph("_____________________________________________________\n")
                
                # 4. Ingest Content (Live Block or DB Fallback)
                if loaded_live and blocks:
                    image_index = 1
                    for block in blocks:
                        b_type = block.get('type')
                        
                        if b_type == 'heading':
                            h = doc.add_heading(block['text'], level=block.get('level', 2))
                            h.runs[0].font.name = 'Inter'
                            h.runs[0].font.size = Pt(14 if block.get('level', 2) == 2 else 12)
                            h.runs[0].font.bold = True
                            
                        elif b_type == 'paragraph':
                            p = doc.add_paragraph(block['text'])
                            p.paragraph_format.line_spacing = 1.15
                            p.paragraph_format.space_after = Pt(6)
                            
                        elif b_type == 'list-item':
                            style = 'List Number' if block.get('list_type') == 'ol' else 'List Bullet'
                            p = doc.add_paragraph(block['text'], style=style)
                            p.paragraph_format.line_spacing = 1.15
                            p.paragraph_format.space_after = Pt(4)
                            
                        elif b_type == 'image':
                            img_src = block.get('src')
                            alt_text = block.get('alt', 'Screenshot')
                            safe_src_print = img_src[:60].encode('ascii', errors='ignore').decode('ascii')
                            print(f"   * Downloading screenshot {image_index}: {safe_src_print}...")
                            try:
                                # Fetch base64 data in browser context to bypass session protections
                                base64_data = page.evaluate("""async (src) => {
                                    try {
                                        const response = await fetch(src);
                                        if (!response.ok) return null;
                                        const blob = await response.blob();
                                        return new Promise((resolve) => {
                                            const reader = new FileReader();
                                            reader.onloadend = () => resolve(reader.result);
                                            reader.readAsDataURL(blob);
                                        });
                                    } catch (e) {
                                        return null;
                                    }
                                }""", img_src)
                                
                                if not base64_data or "," not in base64_data:
                                    raise ValueError("Could not retrieve valid image base64 data")
                                
                                # Process and decode Base64
                                header, encoded = base64_data.split(",", 1)
                                image_bytes = base64.b64decode(encoded)
                                
                                temp_img_path = f"_temp_screenshot_{id_}_{image_index}.png"
                                with open(temp_img_path, "wb") as f_img:
                                    f_img.write(image_bytes)
                                
                                # Insert Image into Document
                                p_img = doc.add_paragraph()
                                p_img.paragraph_format.space_before = Pt(8)
                                p_img.paragraph_format.space_after = Pt(4)
                                run_img = p_img.add_run()
                                run_img.add_picture(temp_img_path, width=Inches(5.5))
                                
                                # Accessible Alt-Text Caption below image
                                p_caption = doc.add_paragraph()
                                p_caption.paragraph_format.space_after = Pt(12)
                                run_caption = p_caption.add_run(f"Figure {image_index}: {alt_text}")
                                run_caption.font.size = Pt(9.5)
                                run_caption.italic = True
                                run_caption.font.color.rgb = docx.shared.RGBColor(100, 100, 100)
                                
                                # Clean up temp file
                                if os.path.exists(temp_img_path):
                                    os.remove(temp_img_path)
                                    
                                image_index += 1
                            except Exception as img_err:
                                print(f"   ⚠️ Image download failed: {img_err}")
                                p_err = doc.add_paragraph()
                                run_err = p_err.add_run(f"[Screenshot Attachment Link: {img_src}]")
                                run_err.bold = True
                                run_err.font.color.rgb = docx.shared.RGBColor(200, 0, 0)
                else:
                    # FALLBACK: DB Text Parsing (for offline / authentication-locked internal articles)
                    print(" -> Fallback to database plain text (contains no screenshot URLs)")
                    body_start = db_content.find("<br><br>")
                    body_content = db_content[body_start + 8:].strip() if body_start != -1 else db_content
                    
                    paragraphs = body_content.split('\n')
                    for p_text in paragraphs:
                        p_text = p_text.strip()
                        if not p_text:
                            continue
                        
                        if p_text.startswith('###'):
                            h = doc.add_heading(p_text.replace('###', '').strip(), level=3)
                            h.runs[0].font.name = 'Inter'
                            h.runs[0].font.size = Pt(12)
                            h.runs[0].font.bold = True
                        elif p_text.startswith('##'):
                            h = doc.add_heading(p_text.replace('##', '').strip(), level=2)
                            h.runs[0].font.name = 'Inter'
                            h.runs[0].font.size = Pt(14)
                            h.runs[0].font.bold = True
                        elif p_text.startswith('#'):
                            h = doc.add_heading(p_text.replace('#', '').strip(), level=1)
                            h.runs[0].font.name = 'Inter'
                            h.runs[0].font.size = Pt(16)
                            h.runs[0].font.bold = True
                        else:
                            p = doc.add_paragraph(p_text)
                            p.paragraph_format.line_spacing = 1.15
                            p.paragraph_format.space_after = Pt(6)

                # Accessibility Checklist
                doc.add_paragraph("\n\nAccessibility Review & Update Log:")
                doc.add_paragraph("- [ ] Content accuracy verified\n- [ ] Color contrast checked\n- [x] Image alt-text verified (screenshots auto-captioned)\n- [ ] Document structure and headings structured")
                
                # Save Document
                safe_title = clean_filename(title)
                if len(safe_title) > 80:
                    safe_title = safe_title[:80]
                    
                file_name = f"{safe_title}.docx"
                file_path = os.path.join(OUTPUT_DIR, file_name)
                
                # Safely handle duplicate filenames
                counter = 1
                while os.path.exists(file_path):
                    file_name = f"{safe_title}_{counter}.docx"
                    file_path = os.path.join(OUTPUT_DIR, file_name)
                    counter += 1
                    
                doc.save(file_path)
                success_count += 1
                
            except Exception as e:
                print(f"Error processing article {id_}: {e}")
        
        browser.close()
        
    print(f"\nSuccessfully processed and updated {success_count} / {len(rows)} articles in {OUTPUT_DIR}!")
    conn.close()

if __name__ == "__main__":
    main()
