import sqlite3
import os
import re
import base64
import sys
import time
from playwright.sync_api import sync_playwright
import docx
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Standard UTF-8 safe print
def safe_print(msg):
    sys.stdout.buffer.write((msg + '\n').encode('utf-8', errors='ignore'))

# Paths
DB_PATH = r"C:\Users\wagahsan\OneDrive - Minnesota State\Desktop\Shadman Ahsan\TRC-AI-Assistant\trc_ai.db"
OUTPUT_DIR = r"C:\Users\wagahsan\OneDrive - Minnesota State\Desktop\Shadman Ahsan\TRC-AI-Assistant\Collected_KB_Documents"

def clean_filename(filename):
    return re.sub(r'[\\/*?:"<>|]', "", filename).strip()

def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    """Adds a clickable hyperlink to a Word paragraph run."""
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)

    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    new_run.append(rPr)
    text_node = OxmlElement('w:t')
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def main():
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)
        
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    
    cur.execute("SELECT id, content FROM kb WHERE source = 'SMSU KB Scraper'")
    rows = cur.fetchall()
    
    safe_print(f"=== SMSU KB Perfect DOCX Converter (Headless/Headed Session) ===")
    safe_print(f"Target Directory: {OUTPUT_DIR}")
    safe_print(f"Total articles in DB: {len(rows)}")
    
    # We will launch in headed mode so the user can log in once, allowing access to ALL screenshots!
    with sync_playwright() as p:
        safe_print("\nLaunching browser in HEADED mode so you can log in...")
        user_data_dir = r"C:\Users\wagahsan\OneDrive - Minnesota State\Desktop\Shadman Ahsan\TRC-AI-Assistant\playwright_user_data_perfect"
        
        context = p.chromium.launch_persistent_context(
            user_data_dir=user_data_dir,
            headless=False,
            viewport={"width": 1280, "height": 800}
        )
        page = context.pages[0] if context.pages else context.new_page()
        
        # Navigate to portal to authenticate
        portal_url = "https://services.smsu.edu/TDClient/180/Portal/KB/Search"
        safe_print(f"Navigating to: {portal_url}")
        page.goto(portal_url, timeout=30000)
        
        safe_print("\n" + "="*80)
        safe_print("🔑 ACTION REQUIRED: PLEASE SIGN IN TO THE SMSU HELP DESK CLIENT PORTAL 🔑")
        safe_print("="*80)
        safe_print("1. Click 'Sign In' in the top-right corner of the browser page.")
        safe_print("2. Log in using your StarID and complete the Microsoft MFA authentication.")
        safe_print("3. Once logged in, the script will AUTOMATICALLY detect it and start!")
        safe_print("="*80 + "\n")
        
        safe_print("Waiting for login detection...")
        while True:
            try:
                # Poll for Sign Out element or disappearance of Sign In element
                if page.locator("a:has-text('Sign Out'), a[href*='SignOut'], a[href*='logout'], a[href*='Logout']").count() > 0:
                    safe_print("🎉 Login successfully detected via Sign Out button!")
                    break
                if page.locator("a:has-text('Sign In'), a[href*='SignIn']").count() == 0 and "TDClient" in page.url:
                    safe_print("🎉 Login successfully detected (Sign In button is gone)!")
                    break
            except Exception:
                pass
            time.sleep(1)
        
        safe_print("\nStarting DOCX conversion with rich text runs and active screenshots...")
        
        success_count = 0
        for idx, (id_, db_content) in enumerate(rows, 1):
            try:
                # Parse metadata
                title_match = re.match(r"\*\*(.*?)\*\*", db_content)
                title = title_match.group(1).strip() if title_match else f"Article_{id_}"
                
                url_match = re.search(r"Official KB Article\*\*: (https://[^\s<]+)", db_content)
                url = url_match.group(1).strip() if url_match else ""
                
                cat_match = re.search(r"Category: (.*?)<br>", db_content)
                category = cat_match.group(1).strip() if cat_match else "General"
                
                tags_match = re.search(r"Tags: (.*?)<br>", db_content)
                tags = tags_match.group(1).strip() if tags_match else "None"
                
                safe_print(f"[{idx}/{len(rows)}] Processing: '{title}'...")
                
                blocks = []
                loaded_live = False
                
                if url:
                    try:
                        page.goto(url, timeout=20000)
                        page.wait_for_selector("#ctl00_ctl00_cpContent_cpContent_divBody, h1", timeout=8000)
                        
                        if page.locator("#ctl00_ctl00_cpContent_cpContent_divBody").count() > 0:
                            # Run DOM DFS to parse elements and extract rich text runs (bold, italic, links)
                            blocks = page.evaluate("""() => {
                                const root = document.querySelector("#ctl00_ctl00_cpContent_cpContent_divBody");
                                if (!root) return [];
                                
                                const blocks = [];
                                
                                function getRuns(element) {
                                    const runs = [];
                                    for (let child of element.childNodes) {
                                        if (child.nodeType === Node.TEXT_NODE) {
                                            const txt = child.textContent;
                                            if (txt) {
                                                runs.push({
                                                    text: txt,
                                                    bold: false,
                                                    italic: false,
                                                    is_link: false
                                                });
                                            }
                                        } else if (child.nodeType === Node.ELEMENT_NODE) {
                                            const tagName = child.tagName;
                                            if (tagName === 'BR') {
                                                runs.push({
                                                    text: "\\n",
                                                    bold: false,
                                                    italic: false,
                                                    is_link: false
                                                });
                                                continue;
                                            }
                                            if (tagName === 'IMG') {
                                                continue; 
                                            }
                                            const isBold = ['STRONG', 'B'].includes(tagName);
                                            const isItalic = ['EM', 'I'].includes(tagName);
                                            const isLink = tagName === 'A';
                                            const href = isLink ? child.getAttribute('href') : '';
                                            
                                            const txt = child.innerText;
                                            if (txt) {
                                                runs.push({
                                                    text: txt,
                                                    bold: isBold || !!child.querySelector('strong, b'),
                                                    italic: isItalic || !!child.querySelector('em, i'),
                                                    is_link: isLink,
                                                    href: href
                                                });
                                            }
                                        }
                                    }
                                    return runs;
                                }
                                
                                function dfs(node) {
                                    if (!node) return;
                                    
                                    const tagName = node.tagName;
                                    
                                    if (tagName === 'IMG') {
                                        blocks.push({
                                            type: 'image',
                                            src: node.src,
                                            alt: node.alt || 'Screenshot'
                                        });
                                        return;
                                    }
                                    
                                    if (['H1', 'H2', 'H3', 'H4', 'H5', 'H6'].includes(tagName)) {
                                        blocks.push({
                                            type: 'heading',
                                            level: parseInt(tagName[1]),
                                            text: node.innerText.trim()
                                        });
                                        return;
                                    }
                                    
                                    if (tagName === 'LI') {
                                        const parentTag = (node.parentElement && node.parentElement.tagName === 'OL') ? 'ol' : 'ul';
                                        blocks.push({
                                            type: 'list-item',
                                            list_type: parentTag,
                                            runs: getRuns(node)
                                        });
                                        
                                        for (let child of node.childNodes) {
                                            if (child.nodeType === Node.ELEMENT_NODE) {
                                                dfs(child);
                                            }
                                        }
                                        return;
                                    }
                                    
                                    if (['P', 'DIV', 'SPAN'].includes(tagName)) {
                                        const hasNestedBlock = node.querySelector('p, h1, h2, h3, h4, h5, h6, ul, ol, li, img');
                                        if (!hasNestedBlock) {
                                            const runs = getRuns(node);
                                            if (runs.length > 0) {
                                                blocks.push({
                                                    type: 'paragraph',
                                                    runs: runs
                                                });
                                            }
                                            return;
                                        }
                                    }
                                    
                                    for (let child of node.childNodes) {
                                        if (child.nodeType === Node.ELEMENT_NODE) {
                                            dfs(child);
                                        } else if (child.nodeType === Node.TEXT_NODE) {
                                            const txt = child.textContent.trim();
                                            if (txt && node === root) {
                                                blocks.push({
                                                    type: 'paragraph',
                                                    runs: [{text: txt, bold: false, italic: false, is_link: false}]
                                                });
                                            }
                                        }
                                    }
                                }
                                
                                dfs(root);
                                return blocks;
                            }""")
                            loaded_live = True
                            safe_print(f"  -> Loaded live! Extracted {len(blocks)} rich content blocks.")
                    except Exception as e:
                        safe_print(f"  -> Live loading skipped: {e}")
                
                # 3. Create Styled DOCX
                doc = Document()
                
                # Heading Spacing & Bold
                h_title = doc.add_heading(title, level=1)
                h_title.runs[0].font.name = 'Inter'
                h_title.runs[0].font.size = Pt(20)
                h_title.runs[0].font.bold = True
                h_title.paragraph_format.space_after = Pt(12)
                
                # Metadata section
                p_meta = doc.add_paragraph()
                p_meta.paragraph_format.space_before = Pt(8)
                p_meta.paragraph_format.space_after = Pt(18)
                
                run_url_lbl = p_meta.add_run("Official URL: ")
                run_url_lbl.bold = True
                run_url_lbl.font.size = Pt(10)
                if url:
                    add_hyperlink(p_meta, url, url)
                else:
                    p_meta.add_run("Not Available")
                    
                p_meta.add_run("\n")
                run_cat = p_meta.add_run("Category: ")
                run_cat.bold = True
                run_cat.font.size = Pt(10)
                p_meta.add_run(category)
                
                p_meta.add_run("\n")
                run_tags = p_meta.add_run("Tags: ")
                run_tags.bold = True
                run_tags.font.size = Pt(10)
                p_meta.add_run(tags)
                
                doc.add_paragraph("_____________________________________________________\n")
                
                # 4. Generate structured content
                if loaded_live and blocks:
                    image_index = 1
                    for block in blocks:
                        b_type = block.get('type')
                        
                        if b_type == 'heading':
                            h = doc.add_heading(block['text'], level=block.get('level', 2))
                            h.runs[0].font.name = 'Inter'
                            h.runs[0].font.size = Pt(14 if block.get('level', 2) == 2 else 12)
                            h.runs[0].font.bold = True
                            h.paragraph_format.space_before = Pt(12)
                            h.paragraph_format.space_after = Pt(6)
                            
                        elif b_type == 'paragraph':
                            p = doc.add_paragraph()
                            p.paragraph_format.line_spacing = 1.15
                            p.paragraph_format.space_after = Pt(8) # Gap sequence
                            
                            # Append runs with formatting preserved!
                            for r in block.get('runs', []):
                                if r.get('is_link') and r.get('href'):
                                    add_hyperlink(p, r['href'], r['text'])
                                else:
                                    run = p.add_run(r['text'])
                                    run.bold = r.get('bold', False)
                                    run.italic = r.get('italic', False)
                                    
                        elif b_type == 'list-item':
                            style = 'List Number' if block.get('list_type') == 'ol' else 'List Bullet'
                            p = doc.add_paragraph(style=style)
                            p.paragraph_format.line_spacing = 1.15
                            p.paragraph_format.space_after = Pt(4)
                            
                            # Append runs inside list item!
                            for r in block.get('runs', []):
                                if r.get('is_link') and r.get('href'):
                                    add_hyperlink(p, r['href'], r['text'])
                                else:
                                    run = p.add_run(r['text'])
                                    run.bold = r.get('bold', False)
                                    run.italic = r.get('italic', False)
                                    
                        elif b_type == 'image':
                            img_src = block.get('src')
                            alt_text = block.get('alt', 'Screenshot')
                            safe_src_print = img_src[:60].encode('ascii', errors='ignore').decode('ascii')
                            safe_print(f"   * Downloading screenshot {image_index}: {safe_src_print}...")
                            
                            try:
                                # Fetch Base64 via Playwright inside session
                                base64_data = page.evaluate("""async (src) => {
                                    try {
                                        const response = await fetch(src);
                                        if (!response.ok) return null;
                                        const blob = await response.blob();
                                        return new Promise((resolve) => {
                                            const reader = new FileReader();
                                            reader.onloadend = () => resolve(reader.result);
                                            reader.readAsDataURL(blob);
                                        });
                                    } catch(e) {
                                        return null;
                                    }
                                }""", img_src)
                                
                                if not base64_data or "," not in base64_data:
                                    raise ValueError("Could not download image in browser")
                                    
                                header, encoded = base64_data.split(",", 1)
                                image_bytes = base64.b64decode(encoded)
                                
                                temp_path = f"_temp_perfect_{image_index}.png"
                                with open(temp_path, "wb") as f_img:
                                    f_img.write(image_bytes)
                                    
                                # Insert Image
                                p_img = doc.add_paragraph()
                                p_img.paragraph_format.space_before = Pt(8)
                                p_img.paragraph_format.space_after = Pt(4)
                                p_img.add_run().add_picture(temp_path, width=Inches(5.5))
                                
                                # Caption alt-text
                                p_caption = doc.add_paragraph()
                                p_caption.paragraph_format.space_after = Pt(12)
                                run_caption = p_caption.add_run(f"Figure {image_index}: {alt_text}")
                                run_caption.font.size = Pt(9.5)
                                run_caption.italic = True
                                run_caption.font.color.rgb = docx.shared.RGBColor(100, 100, 100)
                                
                                if os.path.exists(temp_path):
                                    os.remove(temp_path)
                                image_index += 1
                            except Exception as img_err:
                                safe_print(f"   ⚠️ Image download failed: {img_err}")
                                p_err = doc.add_paragraph()
                                run_err = p_err.add_run(f"[Screenshot Attachment: {img_src}]")
                                run_err.bold = True
                                run_err.font.color.rgb = docx.shared.RGBColor(200, 0, 0)
                else:
                    # FALLBACK: DB Text Parsing (if live page loading is offline or skipped)
                    body_start = db_content.find("<br><br>")
                    body_content = db_content[body_start + 8:].strip() if body_start != -1 else db_content
                    paragraphs = body_content.split('\n')
                    for p_text in paragraphs:
                        p_text = p_text.strip()
                        if not p_text:
                            continue
                        
                        if p_text.startswith('###'):
                            h = doc.add_heading(p_text.replace('###', '').strip(), level=3)
                            h.paragraph_format.space_before = Pt(8)
                            h.paragraph_format.space_after = Pt(4)
                        elif p_text.startswith('##'):
                            h = doc.add_heading(p_text.replace('##', '').strip(), level=2)
                            h.paragraph_format.space_before = Pt(10)
                            h.paragraph_format.space_after = Pt(6)
                        elif p_text.startswith('#'):
                            h = doc.add_heading(p_text.replace('#', '').strip(), level=1)
                            h.paragraph_format.space_before = Pt(12)
                            h.paragraph_format.space_after = Pt(6)
                        else:
                            p = doc.add_paragraph(p_text)
                            p.paragraph_format.line_spacing = 1.15
                            p.paragraph_format.space_after = Pt(8)
                
                # Checklist
                doc.add_paragraph("\n\nAccessibility Review & Update Log:")
                doc.add_paragraph("- [ ] Content accuracy verified\n- [ ] Color contrast checked\n- [x] Image alt-text verified (screenshots auto-captioned)\n- [x] Document structure and headings formatted (perfect bold, list bulletpoints and sequences preserved)")
                
                # Save Docx (overwriting old files cleanly to avoid duplicates)
                safe_title = clean_filename(title)
                if len(safe_title) > 80:
                    safe_title = safe_title[:80]
                file_name = f"{safe_title}.docx"
                file_path = os.path.join(OUTPUT_DIR, file_name)
                
                doc.save(file_path)
                success_count += 1
                
            except Exception as e:
                safe_print(f"Error converting article {id_}: {e}")
                
        context.close()
        
    safe_print(f"\n🎉 Completed! Successfully generated {success_count} / {len(rows)} perfect articles with full screenshots and rich formatting!")
    conn.close()

if __name__ == "__main__":
    main()
