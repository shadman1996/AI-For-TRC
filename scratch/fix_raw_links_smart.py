import os
import re
import sys
import docx
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def safe_print(msg):
    sys.stdout.buffer.write((msg + '\n').encode('utf-8', errors='ignore'))

def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
    new_run.append(rPr)
    text_node = OxmlElement('w:t')
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def get_contextual_word(url):
    """Resolves a clean, user-friendly anchor word based on the URL path."""
    url_lower = url.lower()
    
    if "d2l" in url_lower or "brightspace" in url_lower:
        return "D2L Brightspace Portal"
    elif "zoom" in url_lower:
        return "Zoom Web Portal"
    elif "airtame" in url_lower:
        return "Airtame Wireless Display Setup"
    elif "newform" in url_lower or "ticketrequests" in url_lower:
        return "TeamDynamix Request Form"
    elif "articledet" in url_lower or "portal/kb" in url_lower:
        return "Official Help Desk Article"
    elif "sccm" in url_lower:
        return "SCCM Software Center"
    elif "printing" in url_lower or "papercut" in url_lower:
        return "SMSU Printing Portal"
    elif "smsu.edu" in url_lower:
        return "SMSU Official Website"
    elif "minnstate.edu" in url_lower:
        return "MinnState Student Portal"
    elif "mailto:" in url_lower:
        email = url.replace("mailto:", "").strip()
        return f"Email {email}"
    
    return "Click Here to Open Link"

def is_copy_paste_url(url, pre_text):
    """Checks if the URL is a technical address/feed that the user must see or copy/paste."""
    url_lower = url.lower()
    pre_lower = pre_text.lower()
    return (
        re.search(r'\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}', url) or 
        url_lower.endswith(".aspx") or 
        url_lower.endswith(".feed") or 
        "feed" in url_lower or 
        "vlan" in url_lower or
        "ip" in url_lower or
        any(word in pre_lower for word in ["copy", "paste", "enter", "workspace url", "feed url", "server", "address", "config"])
    )

def smart_replace_links(text):
    """Parses text and extracts raw URLs, converting them into clean segments."""
    url_pattern = re.compile(r'(https?://[^\s()<>]+)')
    matches = list(url_pattern.finditer(text))
    if not matches:
        return [('text', text)]
        
    segments = []
    last_idx = 0
    
    # CASE A: Explicit Label Before URL (e.g. "Label: URL" or "Label - URL" or "Label: [URL]")
    # If the paragraph contains exactly one URL, and has a colon or hyphen, and the label is reasonable:
    if len(matches) == 1:
        match = matches[0]
        start = match.start()
        orig_end = match.end()
        url = match.group(0)
        
        # Clean trailing punctuation from URL
        trailing_punc = ""
        while url and url[-1] in [".", ",", ";", ")", "]", ":"]:
            trailing_punc = url[-1] + trailing_punc
            url = url[:-1]
            
        pre_text = text[:start].strip()
        post_text = text[orig_end:].strip()
        
        # If it is a copy-paste URL, do NOT treat it as a pure label replacement
        if is_copy_paste_url(url, pre_text):
            pass
        elif pre_text.endswith(":") or pre_text.endswith("-"):
            label = pre_text[:-1].strip()
            # If the label is short and descriptive (not a full long sentence)
            if 3 < len(label) < 60 and not any(verb in label.lower() for verb in ["click", "submit", "open", "please", "go to", "visit"]):
                segments.append(('link', url, label))
                if trailing_punc:
                    segments.append(('text', trailing_punc))
                if post_text:
                    segments.append(('text', " " + post_text))
                return segments

    # Otherwise, parse sequentially
    for match in matches:
        start = match.start()
        orig_end = match.end()
        url = match.group(0)
        
        # Clean trailing punctuation from URL
        trailing_punc = ""
        while url and url[-1] in [".", ",", ";", ")", "]", ":"]:
            trailing_punc = url[-1] + trailing_punc
            url = url[:-1]
            
        pre_text = text[last_idx:start]
        
        if is_copy_paste_url(url, pre_text):
            # Keep the raw URL visible and hyperlinked to itself
            if pre_text:
                segments.append(('text', pre_text))
            segments.append(('link', url, url))
            if trailing_punc:
                segments.append(('text', trailing_punc))
        else:
            # CASE C: Wrapped in brackets/parentheses, e.g. "Go to the Mobility Print app [http://...]"
            is_bracketed = False
            bracket_char = ""
            if pre_text and pre_text[-1] in ["[", "("]:
                bracket_char = pre_text[-1]
                is_bracketed = True
                
            if is_bracketed:
                # Remove the bracket from pre_text
                pre_text_clean = pre_text[:-1].rstrip()
                # Find a good anchor phrase in pre_text_clean
                words = pre_text_clean.split()
                if len(words) >= 3:
                    anchor_text = " ".join(words[-3:])
                    pre_text_base = " ".join(words[:-3]) + " "
                elif len(words) > 0:
                    anchor_text = " ".join(words)
                    pre_text_base = ""
                else:
                    anchor_text = get_contextual_word(url)
                    pre_text_base = ""
                    
                if pre_text_base:
                    segments.append(('text', pre_text_base))
                segments.append(('link', url, anchor_text))
                
                # Check if there is a closing bracket in trailing_punc
                closing_bracket = "]" if bracket_char == "[" else ")"
                if trailing_punc and trailing_punc.startswith(closing_bracket):
                    trailing_punc = trailing_punc[1:]
                if trailing_punc:
                    segments.append(('text', trailing_punc))
            else:
                # CASE D: Standalone URL replacement
                if pre_text:
                    segments.append(('text', pre_text))
                anchor_text = get_contextual_word(url)
                segments.append(('link', url, anchor_text))
                if trailing_punc:
                    segments.append(('text', trailing_punc))
                    
        last_idx = orig_end
        
    if last_idx < len(text):
        segments.append(('text', text[last_idx:]))
        
    return segments

def has_clean_hyperlinks(paragraph):
    """Checks if the paragraph contains hyperlinks that already have clean contextual display text (not raw URLs)."""
    if 'w:hyperlink' in paragraph._p.xml:
        hyperlinks = paragraph._p.xpath('.//w:hyperlink')
        for hl in hyperlinks:
            hl_text = "".join(hl.xpath('.//w:t/text()')).strip()
            # If the hyperlink displays a clean name instead of a raw URL, it is clean!
            if hl_text and not re.match(r'^https?://[^\s()<>]+$', hl_text):
                return True
    return False

def clear_paragraph(paragraph):
    """Safely and completely clears all XML children (runs, hyperlinks) of a paragraph while preserving formatting/style."""
    p_elem = paragraph._p
    pPr = p_elem.pPr
    p_elem.clear()
    if pPr is not None:
        p_elem.append(pPr)

def process_paragraph_links(paragraph):
    """Finds raw URLs (including raw-URL hyperlinks) and converts them to beautifully styled hyperlinks."""
    text = paragraph.text
    url_pattern = re.compile(r'(https?://[^\s()<>]+)')
    matches = list(url_pattern.finditer(text))
    
    if not matches:
        return False
        
    # Guard: Skip only if the paragraph has a clean contextual hyperlink already
    if has_clean_hyperlinks(paragraph):
        return False
        
    segments = smart_replace_links(text)
    
    # Safe and complete XML clear of paragraph contents (clears old ghost hyperlinks)
    clear_paragraph(paragraph)
    
    for seg_type, *seg_data in segments:
        if seg_type == 'text':
            content = seg_data[0]
            run = paragraph.add_run(content)
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
        elif seg_type == 'link':
            url, anchor = seg_data
            add_hyperlink(paragraph, url, anchor)
            
    return True

def fix_links_in_docx(file_path):
    try:
        doc = Document(file_path)
        modified = False
        
        # 1. Process standard body paragraphs
        for p in doc.paragraphs:
            # Skip metadata key/value paragraphs
            if p.text.strip().startswith("Official URL:") or p.text.strip().startswith("Source URL:") or p.text.strip().startswith("Request Forms / URL"):
                continue
            if process_paragraph_links(p):
                modified = True
                
        # 2. Process table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Skip metadata table rows
                    if "https://" in cell.text or "http://" in cell.text:
                        if "Request Forms" in cell.text or "Source URL" in cell.text:
                            continue
                        
                        # Guard: Check cell paragraph XML for clean hyperlinks
                        has_clean = False
                        for p in cell.paragraphs:
                            if has_clean_hyperlinks(p):
                                has_clean = True
                                break
                        if has_clean:
                            continue
                            
                        for p in cell.paragraphs:
                            if process_paragraph_links(p):
                                modified = True
                                
        if modified:
            doc.save(file_path)
            return True
    except Exception as e:
        safe_print(f"  ⚠️ Error processing {os.path.basename(file_path)}: {e}")
    return False

def main():
    target_dirs = [
        r"c:\Users\wagahsan\OneDrive - Minnesota State\Desktop\Shadman Ahsan\TRC-AI-Assistant\Revised_KB_Documents",
        r"c:\Users\wagahsan\OneDrive - Minnesota State\Desktop\Shadman Ahsan\TRC-AI-Assistant\Supervisor_Target_Articles",
        r"c:\Users\wagahsan\OneDrive - Minnesota State\Desktop\Shadman Ahsan\TRC-AI-Assistant\Revised_Service_Catalog_Documents",
        r"c:\Users\wagahsan\OneDrive - Minnesota State\Desktop\Shadman Ahsan\TRC-AI-Assistant\Supervisor_Target_Service_Catalog"
    ]
    
    safe_print("=== Smart Raw Links to Word Hyperlinks Converter ===")
    
    fixed_count = 0
    total_scanned = 0
    
    for base_dir in target_dirs:
        if not os.path.exists(base_dir):
            continue
        safe_print(f"\nScanning Folder: {base_dir}")
        for root, dirs, files in os.walk(base_dir):
            for file in files:
                if file.endswith(".docx"):
                    file_path = os.path.join(root, file)
                    total_scanned += 1
                    if fix_links_in_docx(file_path):
                        fixed_count += 1
                        safe_print(f"  🔗 Smart-Hyperlinked Raw URLs in: {os.path.relpath(file_path, base_dir)}")
                        
    safe_print(f"\n=== Conversion Finished ===")
    safe_print(f"Scanned {total_scanned} total documents.")
    safe_print(f"Successfully converted raw plain-text URLs to styled hyperlinks in {fixed_count} files!")

if __name__ == "__main__":
    main()
